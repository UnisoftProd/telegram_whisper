import telebot
from telebot import apihelper
import subprocess
import os, logging
from telebot.types import LabeledPrice, ShippingOption
import  psycopg2, psycopg2.extras
from sqlalchemy import create_engine, Column, Integer, String, select, update, exists, Boolean
from sqlalchemy.orm import sessionmaker
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.orm import declarative_base
from docx import Document
from telebot import types
import telebot.apihelper
import re
from difflib import get_close_matches
from threading import Thread
import threading
import yandex_downloader
import time
import gdown
import streamyard_downloader
import mutagen 
from mutagen.wave import WAVE 
import shutil
from zipfile import ZipFile
import short_descriptor
import locale
import _locale

VIDEO_EXTENSIONS = {".mp4", ".avi", ".mkv", ".mov", ".wmv", ".flv"}
AUDIO_EXTENSIONS = {".mp3", ".wav", ".aac", ".flac", ".ogg", ".wma"}
media_groups_in_work = {}

#bc5faeb33eb16909b64fdc3bb397e3aa
#25730961
#./telegram-bot-api.exe --local --api-id=25730961 --api-hash=bc5faeb33eb16909b64fdc3bb397e3aa
"""
CREATE TABLE users(
id BIGSERIAL NOT NULL PRIMARY KEY,
user_tg BIGSERIAL NOT NULL,
paid BOOLEAN NOT NULL,
minutes INTEGER
);

"""
Base = declarative_base()

url_pattern = re.compile(
        r'^(https?://)?(www\.)?([a-zA-Z0-9_\-]+(\.[a-zA-Z]{2,})+)(/[^\s]*)?$'
    )

class Users(Base):
    __tablename__ = 'users'
    id = Column(Integer, primary_key=True, autoincrement=True)
    user_tg = Column(Integer, nullable=False)
    paid = Column(Boolean, nullable=False)
    minutes = Column(Integer)


class PostgresManager:
    def __init__(self, connection_url):
        try:
            self.engine = create_engine(connection_url)
            self.Session = sessionmaker(bind=self.engine)
            Base.metadata.create_all(self.engine)
        except SQLAlchemyError as e:
            logging.error(f"Ошибка подключения к базе данных: {e}")
            raise

    def make_new_session(self):
        self.session = self.Session()

    def close_session(self):
        self.session.close()

    def is_in_database(self, user_id):
        try:
            user = self.session.query(exists().where(Users.user_tg == user_id)).scalar()
            if user:
                return True
            else:
                return False
                
        except SQLAlchemyError as e:
            logging.error(f"Ошибка проверки существования записи: {e}")
            self.session.close()
            raise

    def is_paid_user(self, user_id):
        try:
            user = self.session.query(Users).filter(Users.user_tg == user_id).one_or_none()
            return user.paid
                
        except SQLAlchemyError as e:
            logging.error(f"Ошибка проверки существования записи: {e}")
            self.session.close()
            raise

    def add_user(self, user_id):
        try:
            new_user = Users(
                user_tg = user_id,
                paid = False
            )
            self.session.add(new_user)
            self.session.commit()
        except SQLAlchemyError as e:
            logging.error(f"Ошибка добавления пользователя: {e}")
            self.session.rollback()
            self.session.close()
            raise

    def get_user(self, user_id):
        try:
            user = self.session.query(Users).filter(Users.user_tg == user_id).one_or_none()
            return user
                
        except SQLAlchemyError as e:
            logging.error(f"Ошибка проверки существования записи: {e}")
            self.session.close()
            raise
    
    def set_new_minutes(self, user_id, new_minutes):
        try:
            self.session.query(Users).filter(Users.user_tg == user_id).update(
                {Users.paid: True, Users.minutes:new_minutes}
            )
            self.session.commit()
        except SQLAlchemyError as e:
            logging.error(f"Ошибка обновления минут: {e}")
            self.session.rollback()
            self.session.close()
            raise       

def delete_media_files(folder_path):
    if not os.path.exists(folder_path):
        print(f"Папка {folder_path} не существует.")
        return
    
    files_deleted = 0

    for root, _, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            file_extension = os.path.splitext(file)[1].lower()
            
            if file_extension in VIDEO_EXTENSIONS or file_extension in AUDIO_EXTENSIONS:
                try:
                    os.remove(file_path)
                    print(f"Удален файл: {file_path}")
                    files_deleted += 1
                except Exception as e:
                    print(f"Ошибка при удалении файла {file_path}: {e}")

    print(f"Удалено {files_deleted} файлов.")

def delete_folders_with_substring(directory, substring):
    try:
        # Проверяем, существует ли директория
        if not os.path.exists(directory):
            print(f"Директория {directory} не существует.")
            return

        # Итерируемся по элементам в директории
        for item in os.listdir(directory):
            item_path = os.path.join(directory, item)
            
            # Проверяем, является ли элемент папкой и содержит ли подстроку
            if os.path.isdir(item_path) and substring in item:
                try:
                    shutil.rmtree(item_path)  # Удаляем пустую папку
                    print(f"Удалена папка: {item_path}")
                except OSError:
                    print(f"Папка {item_path} не пуста или произошла ошибка при удалении.")

    except Exception as e:
        print(f"Произошла ошибка: {e}")

def threadwrap(threadfunc):
    def wrapper(*args, **kwargs):
        while True:
            time.sleep(5)
            try:
                threadfunc(*args, **kwargs)
            except BaseException as e:
                print(e)
                th_name = threading.current_thread().name
                print(f'Падение потока {th_name}, перезапуск...')
    return wrapper

@threadwrap
def queue_processor(bot):
    while media_queue:
        if media_queue[0][0].content_type == "voice" or media_queue[0][0].content_type == "video"  or media_queue[0][0].content_type == "audio"  or media_queue[0][0].content_type == "document":
            
            process_audio(bot, media_queue[0][0], "")
        else:
            process_audio(bot, media_queue[0][0], media_queue[0][1])
        media_queue.pop(0)
    time.sleep(15)
    delete_folders_with_substring(os.getcwd(), "google")
    delete_media_files(os.getcwd())
    return


db_manager = PostgresManager("postgresql://postgres:0451@localhost/telegramaudiobot")
bot = telebot.TeleBot('7666679307:AAEvK-aGkdPixr7_8QgnUV7FuC699eIE7kU')

provider_token = '1744374395:TEST:4362222678581d074fc5' 
apihelper.API_URL = "http://localhost:8081/bot{0}/{1}"
apihelper._get_req_session().timeout = (1000, 9000)
media_queue = []
meida_queue_thread = Thread(target=queue_processor, args=(bot,))


prices = [LabeledPrice(label='Working Time Machine', amount=2000)]
start_time_processing = 0.0

# shipping_options = [
#     ShippingOption(id='instant', title='WorldWide Teleporter').add_price(LabeledPrice('Teleporter', 1000)),
#     ShippingOption(id='pickup', title='Local pickup').add_price(LabeledPrice('Pickup', 300))]

users_to_summury = {}



def format_seconds(seconds):
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    seconds = seconds % 60
    return f"{hours:02}:{minutes:02}:{seconds:02}"

def txt_to_word(txt_file_path, word_file_path):
    """
    Конвертирует текстовый файл в документ Word.
    
    :param txt_file_path: Путь к исходному текстовому файлу.
    :param word_file_path: Путь для сохранения выходного файла Word.
    """
    try:
        # Создаем новый документ Word
        doc = Document()
        
        # Читаем содержимое текстового файла
        with open(txt_file_path, 'r', encoding='utf-8') as txt_file:
            lines = txt_file.readlines()
        
        # Добавляем содержимое файла в документ Word
        for line in lines:
            doc.add_paragraph(line.strip())
        
        # Сохраняем документ Word
        doc.save(word_file_path)
        print(f"Файл успешно сохранен как {word_file_path}")
    except Exception as e:
        print(f"Произошла ошибка: {e}")

def rename_file(directory, reference_string):
    # Получить список всех файлов в директории
    files = os.listdir(directory)
    
    # Найти файл, наиболее похожий на строку reference_string
    matches = get_close_matches(reference_string, files, n=1, cutoff=0.1)
    
    if matches:
        original_file = matches[0]
        # Создать новое имя файла
        new_file_name = f"{reference_string}"  # Измените по вашему желанию
        original_path = os.path.join(directory, original_file)
        new_path = os.path.join(directory, new_file_name)
        
        # Переименовать файл
        os.rename(original_path, new_path)
        print(f"Файл {original_file} переименован в {new_file_name}")
        return new_file_name
    else:
        print("Не найден файл, похожий на заданную строку.")
        return reference_string

def insert_in_queue(media_queue, message, result):
    first_id = media_queue[0][0].from_user.id
    self_already_was = False
    for i in range(1, len(media_queue)):
        if media_queue[i][0].from_user.id == message.from_user.id:
            self_already_was = True
        if media_queue[i][0].from_user.id == first_id:
            if self_already_was:
                self_already_was = False
                continue
            else:
                media_queue.insert(i, [message, result])
                return
    media_queue.append([message, result])

def get_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """

    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment

def merge_word_documents(file1, file2, name):
    """
    Объединяет содержимое второго Word-файла в первый, добавляя имя второго файла перед содержимым.

    :param file1: Путь к первому Word-файлу (docx), в который будет добавлено содержимое.
    :param file2: Путь ко второму Word-файлу (docx), содержимое которого будет добавлено.
    """
    # Открываем первый документ
    doc1 = Document(file1)

    # Открываем второй документ
    doc2 = Document(file2)

    # Добавляем имя второго файла в первый документ
    doc1.add_paragraph(f"Содержимое файла: {name}")


    # Добавляем содержимое второго документа в первый
    for para in doc2.paragraphs:
        get_para_data(doc1, para)

    doc1.add_page_break()
    # Сохраняем изменения в первом документе
    doc1.save(file1)

def link_downloader(command, message):
    if "disk.yandex.ru/d" in message.text:
            try:
                yandex_downloader.process_public_folder(message.text, media_queue, message, meida_queue_thread, bot, media_groups_in_work, local_dir="./")
                return
            except Exception as e:
                print(e)
                pass
    elif "drive.google.com/drive/folders" in message.text:
        try:
            os.makedirs("google_" + str(message.id))
            gdown.download_folder(media_groups_in_work, media_queue, message, meida_queue_thread, bot, 
                                  message.text, quiet=False, remaining_ok=True, use_cookies=False, output="google_" + str(message.id))
            return
        except Exception as e:
            print(e)
            pass
    elif "streamyard.com" in message.text:
        try:
            streamyard_downloader.download_streamyard(message.text, message, media_queue, meida_queue_thread)
            return
        except Exception as e:
            print(e)
            pass
    try:
        _locale._getdefaultlocale = (lambda *args: ['en_US', 'cp1252'])
        print(locale.getpreferredencoding())

        result = subprocess.run(command, capture_output=True, text=True)
        print(result)
        # Проверяем успешность выполнения команды
        if result.returncode != 0:
            bot.send_message(message.chat.id, "Ошибка при загрузке файла")
            return None

        # Извлекаем имя файла из вывода команды
        filename = ""
        lines = result.stderr.splitlines() + result.stdout.splitlines()
        for line in lines:
            if "[download] Destination:" in line:
                raw_filename = line.split("[download] Destination:")[1].strip()
                raw_filename = raw_filename.split(".")[0] + ".wav"
                directory = "./"  # Укажите путь к директории
                rename_file(directory, raw_filename)
                filename = raw_filename


        if filename:
            if media_queue:
                insert_in_queue(media_queue, message, filename)
            else:
                media_queue.append([message, filename])
            bot.send_message(message.chat.id, "Файл добавлен в очередь")
            if not meida_queue_thread.is_alive():
                meida_queue_thread.start()
    except Exception:
        bot.send_message(message.chat.id, "Ошибка при загрузке файла")

def process_audio(bot, message, ready_file = ""):
    global start_time_processing
    start_time_processing = time.time()
    if ready_file != "":
        file_audio_name = ready_file
    elif message.content_type == "voice":
        media_file = message.voice
        a = bot.get_file(media_file.file_id)
        process = subprocess.run(['ffmpeg', "-y", '-i', a.file_path, str(a.file_id) + ".wav"])
        file_audio_name = str(a.file_id) + ".wav"
    elif message.content_type == 'video':
        media_file = message.video
        a = bot.get_file(media_file.file_id)
        process = subprocess.run(['ffmpeg', "-y", '-i', a.file_path, '-vn', '-acodec', 'pcm_s16le', '-ar', '44100', '-ac', '2', str(a.file_id) + ".wav"])
        file_audio_name = str(a.file_id) + ".wav"
    elif message.content_type == 'document': 
        media_file = message.document
        a = bot.get_file(media_file.file_id)
        process = subprocess.run(['ffmpeg', "-y", '-i', a.file_path, '-vn', '-acodec', 'pcm_s16le', '-ar', '44100', '-ac', '2', str(a.file_id) + ".wav"])
        file_audio_name = str(a.file_id) + ".wav"
    else:
        media_file = message.audio
        a = bot.get_file(media_file.file_id)
        process = subprocess.run(['ffmpeg', "-y", '-i', a.file_path, str(a.file_id) + ".wav"])
        file_audio_name = str(a.file_id) + ".wav"
    
    bot.send_message(message.chat.id, "Обработка. Пожалуйста, подождите")
    try:
        subprocess.run(["python", "diarize.py", "-a", file_audio_name, "--no-stem", "--whisper-model", "small"])
    except Exception:
        return False
    
    file_audio_name = file_audio_name[:file_audio_name.rfind('.')] if '.' in file_audio_name else file_audio_name

    file_path = f"{file_audio_name}.txt"

    if os.path.exists(file_path):

        if message.from_user.id in users_to_summury.keys() and users_to_summury[message.from_user.id]:
            description = short_descriptor.make_description(file_path)
            bot.send_message(message.chat.id, description)
            print(description)


        word_file_path = f"{file_audio_name}.docx"
        txt_to_word(file_path, word_file_path)


        if message.media_group_id in media_groups_in_work.keys():
            if len(media_groups_in_work[message.media_group_id]) > 1:
                name_word = ""
                if ready_file != "":
                    name_word = os.path.basename(file_audio_name + ".docx")
                elif  message.content_type == "voice":
                    name_word = "Голосовое сообщение" + ".docx"
                else:
                    name_word = media_file.file_name

                with ZipFile(str(message.id) + ".zip", "a") as myzip:
                    myzip.write(word_file_path)

                if os.path.exists(str(message.media_group_id) + ".docx"):
                     merge_word_documents(str(message.media_group_id) + ".docx", word_file_path, name_word)
                else:
                    doc = Document()
                    doc.save(str(message.media_group_id) + ".docx")
                    merge_word_documents(str(message.media_group_id) + ".docx", word_file_path, name_word)
                if media_groups_in_work[message.media_group_id][-1] == message or media_groups_in_work[message.media_group_id][-1] == ready_file:
                    with open(str(message.id) + ".zip", 'rb') as file:
                        bot.send_document(message.chat.id, file, visible_file_name = "Все файлы из группы" + ".zip")
                    with open(str(message.media_group_id) + ".docx", 'rb') as file:
                        bot.send_document(message.chat.id, file, visible_file_name = "Все файлы из группы" + ".docx")

                    os.remove(str(message.id) + ".zip")
                    os.remove(str(message.media_group_id) + ".docx")
                    del media_groups_in_work[message.media_group_id]

        with open(word_file_path, 'rb') as file:
            if ready_file != "":
                new_filename = os.path.basename(file_audio_name + ".docx")
                #bot.send_document(message.chat.id, file, visible_file_name = ready_file.split(".")[0] + ".docx")
                bot.send_document(message.chat.id, file, visible_file_name = new_filename)
            elif  message.content_type == "voice":
                bot.send_document(message.chat.id, file, visible_file_name = "Голосовое сообщение" + ".docx")
            else:
                file_audio_name1 = media_file.file_name[:media_file.file_name.rfind('.')] if '.' in media_file.file_name else media_file.file_name
                bot.send_document(message.chat.id, file, visible_file_name = file_audio_name1 + ".docx")
    else:
        print(f"Файл {file_path} не найден")

    
    file_path = f"{file_audio_name}.srt"
    if os.path.exists(file_path):
        with open(file_path, 'rb') as file:
            if ready_file != "":
                new_filename = os.path.basename(file_audio_name + "_субтитры.srt")
                #bot.send_document(message.chat.id, file, visible_file_name = ready_file.split(".")[0] + "_субтитры.srt")
                bot.send_document(message.chat.id, file, visible_file_name = new_filename)
            elif  message.content_type == "voice":
                bot.send_document(message.chat.id, file, visible_file_name = "Голосовое сообщение" + "_субтитры.srt")
            else:
                #bot.send_document(message.chat.id, file, visible_file_name = media_file.file_name.split(".")[0] + "_субтитры.srt")
                file_audio_name1 = media_file.file_name[:media_file.file_name.rfind('.')] if '.' in media_file.file_name else media_file.file_name
                bot.send_document(message.chat.id, file, visible_file_name = file_audio_name1 + "_субтитры.srt")
    else:
        print(f"Файл {file_path} не найден")

    try:
        os.remove(str(file_audio_name) +".wav")
        os.remove(str(file_audio_name) +".docx")
        os.remove(str(file_audio_name) +".srt")
        os.remove(str(file_audio_name) +".txt") #тут ошибка  
    except Exception:
        print("Ошибка при удалении файла")
    
    return True



@bot.message_handler(commands=['start'])
def handle_start(message):
    db_manager.make_new_session()
    if not db_manager.is_in_database(message.from_user.id):
            db_manager.add_user(message.from_user.id)
    #keyboard = types.InlineKeyboardMarkup()
    #button1 = types.InlineKeyboardButton(text="Купить минуты", callback_data="button1")
    #keyboard.add(button1)
    keyboard = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    button_buy = telebot.types.KeyboardButton(text="Купить минуты")
    button_summmury = telebot.types.KeyboardButton(text="Генерировать краткое содержание")
    if message.from_user.id == 197016529 or message.from_user.id == 5085651430 or message.from_user.id == 220183123:
        button_restart = telebot.types.KeyboardButton(text="Перезапуск бота")
        keyboard.add(button_restart)
    keyboard.add(button_buy)
    keyboard.add(button_summmury)

    bot.send_message(message.chat.id, "Я перевожу голосовые сообщения и аудиофайлы в текст.  Файл продолжительностью до 2 минут - бесплатно.",  reply_markup=keyboard)

    db_manager.close_session()

@bot.message_handler(commands=['balance'])
def handle_start(message):
    db_manager.make_new_session()
    if not db_manager.is_in_database(message.from_user.id):
            db_manager.add_user(message.from_user.id)
    user  = db_manager.get_user(message.from_user.id)
    if user.paid:
        bot.send_message(message.chat.id, "У вас на балансе " + str(int(user.minutes / 60)) + " минут")
    else:
        bot.send_message(message.chat.id, "У вас бесплатный аккаунт, вы можете транскриировать аудио продолжительностью не более 10 минут")
    
    db_manager.close_session()


@bot.message_handler(commands=['check_transcribs'])
def handle_trascribs_list(message):
    keyboard = types.InlineKeyboardMarkup()
    button_cancel = types.InlineKeyboardButton(text="Удалить очередь", callback_data="cancel_button")
    button_stop_transcrib = types.InlineKeyboardButton(text="Прервать транскрибацию", callback_data="stop_transc_button")
    keyboard.add(button_cancel)

    transcrib_string = ""
    count = 1

    full_length = 0
    temp_length = 0
    for i in range(len(media_queue)):
        if media_queue[i][0].content_type == "voice" or media_queue[i][0].content_type == "video"  or media_queue[i][0].content_type == "audio"  or media_queue[i][0].content_type == "document":
            if media_queue[i][0].content_type == "voice":
                media_file = media_queue[i][0].voice
                length = media_file.duration
            elif media_queue[i][0].content_type == 'video':
                media_file = media_queue[i][0].video
                length = media_file.duration
            elif media_queue[i][0].content_type == 'document': #DURATION проверить изменить
                media_file = media_queue[i][0].document
                a = bot.get_file(media_file.file_id)
                a = a.file_path
                audio = WAVE(str(a)) 
                audio_info = audio.info 
                length = int(audio_info.length) 
            elif media_queue[i][0].content_type == 'audio': #DURATION проверить изменить
                media_file = media_queue[i][0].audio
                length = media_file.duration
        else:
            audio = WAVE(str(media_queue[i][1])) 
            audio_info = audio.info 
            length = int(audio_info.length) 

        if start_time_processing and i == 0:
            difference = int(time.time() - start_time_processing)
            length -= difference
            if length < 0:
                length = 0
        temp_length += length


        if media_queue[i][0].from_user.id == message.from_user.id:
            if i == 0:
                keyboard.add(button_stop_transcrib)
            duration = format_seconds(length) 
            str_file = os.path.basename(str(media_queue[i][1]))
            transcrib_string += "Файл " + str(count) + ": " + str_file + "\nПримерное время готовности: " + duration + "\nПозиция в очереди: " + str(i) +"\n\n"
            count += 1
            full_length += temp_length
            temp_length = 0
    if transcrib_string != "":
        duration = format_seconds(full_length) 
        transcrib_string += f"Общее время готовности: {duration}"
        bot.send_message(message.chat.id, transcrib_string, reply_markup=keyboard)
    else:
        bot.send_message(message.chat.id, "У вас нет активных транскрибаций")


@bot.message_handler(func=lambda message: message.text == 'Генерировать краткое содержание')
def generate_summury(message):
    keyboard = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    button_buy = telebot.types.KeyboardButton(text="Купить минуты")
    button_summmury = telebot.types.KeyboardButton(text="Не генерировать краткое содержание")
    if message.from_user.id == 197016529 or message.from_user.id == 5085651430 or message.from_user.id == 220183123:
        button_restart = telebot.types.KeyboardButton(text="Перезапуск бота")
        keyboard.add(button_restart)
    keyboard.add(button_buy)
    keyboard.add(button_summmury)
    users_to_summury[message.from_user.id] = True
    bot.send_message(message.chat.id, "Бот будет высылать краткое содержание, сгенерированное нейросетью",  reply_markup=keyboard)

@bot.message_handler(func=lambda message: message.text == 'Не генерировать краткое содержание')
def generate_summury(message):
    keyboard = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    button_buy = telebot.types.KeyboardButton(text="Купить минуты")
    button_summmury = telebot.types.KeyboardButton(text="Генерировать краткое содержание")
    if message.from_user.id == 197016529 or message.from_user.id == 5085651430 or message.from_user.id == 220183123:
        button_restart = telebot.types.KeyboardButton(text="Перезапуск бота")
        keyboard.add(button_restart)
    keyboard.add(button_buy)
    keyboard.add(button_summmury)
    users_to_summury[message.from_user.id] = False
    bot.send_message(message.chat.id, "Бот не будет высылать краткое содержание, сгенерированное нейросетью",  reply_markup=keyboard)




@bot.message_handler(func=lambda message: message.text == 'Купить минуты')
def buy_minuts(message):
    keyboard = types.InlineKeyboardMarkup()
    button_buy_1 = types.InlineKeyboardButton(text="Купить час", callback_data="button2")
    button_buy_2 = types.InlineKeyboardButton(text="Купить 5 часов", callback_data="button3")
    button_buy_3 = types.InlineKeyboardButton(text="Купить 10 часов", callback_data="button4")
    keyboard.add(button_buy_1)
    keyboard.add(button_buy_2)
    keyboard.add(button_buy_3)
    bot.send_message(message.chat.id, "Выберите тариф",  reply_markup=keyboard)


@bot.message_handler(func=lambda message: message.text == 'Перезапуск бота')
def reboot_bot_handler(message):
    if message.from_user.id == 197016529 or message.from_user.id == 5085651430 or message.from_user.id == 220183123:
        delete_media_files(os.getcwd())
        delete_folders_with_substring(os.getcwd(), "google")
        script_to_run = "c:/good_aufio_to_text/whisper-diarization-main/reboot_bot.py"
        python_executable = "C:/good_aufio_to_text/my_test_venv/Scripts/python.exe"
        subprocess.Popen([python_executable, script_to_run])



@bot.message_handler(func=lambda message: url_pattern.match(message.text) is not None)
def get_eida_by_link(message):

    # Формируем базовую команду для yt-dlp
    command = ["yt-dlp", message.text]
    command.extend(["-x", "--audio-format", "wav", "--cookies-from-browser", "firefox"])
    output_template = "%(title)s.%(ext)s"
    command.extend(["-o", output_template])

    bot.send_message(message.chat.id, f"Начинается загрузка из {message.text}")

    th = Thread(target=link_downloader, args=(command, message,))
    th.start()


@bot.callback_query_handler(func=lambda call: True)
def callback_inline(call):
    if call.message:      
        if call.data == "button2":
            bot.send_message(call.message.chat.id,
                     "Покупка часа")
            bot.send_invoice(
                     call.message.chat.id,  #chat_id
                     'Покупка 1 часа перевода аудио в текст', 
                     'Вам будет доступен перевод аудиофайлов с распознаванием голосов и соблюденем грамматики общей длительностью до 1 часа', 
                     '1 час', 
                     provider_token, 
                     'RUB', #currency
                     [LabeledPrice(label='1 час', amount=50000)], 
                     #photo_url='pic.webp',
                     photo_height=512,  
                     photo_width=512,
                     photo_size=512,
                     is_flexible=False, 
                     )
        
        if call.data == "button3":
            bot.send_message(call.message.chat.id,
                     "Покупка 5 часов")
            bot.send_invoice(
                     call.message.chat.id, 
                     'Покупка 5 часов перевода аудио в текст', 
                     'Вам будет доступен перевод аудиофайлов с распознаванием голосов и соблюденем грамматики общей длительностью до 5 часов',
                     '5 часов', 
                     provider_token, 
                     'RUB', 
                     [LabeledPrice(label='5 часов', amount=100000)], 
                     #photo_url='pic.webp',
                     photo_height=512,  
                     photo_width=512,
                     photo_size=512,
                     is_flexible=False,  
                     )
        
        if call.data == "button4":
            bot.send_message(call.message.chat.id,
                     "Покупка 10 часов")
            bot.send_invoice(
                     call.message.chat.id,  
                     'Покупка 10 часов перевода аудио в текст', 
                     'Вам будет доступен перевод аудиофайлов с распознаванием голосов и соблюденем грамматики общей длительностью до 10 часов', 
                     '10 часов', 
                     provider_token,
                     'RUB',
                     [LabeledPrice(label='10 часов', amount=250000)], 
                     #photo_url='pic.webp',
                     photo_height=512,  
                     photo_width=512,
                     photo_size=512,
                     is_flexible=False, 
                     )
        
        elif call.data == "cancel_button":
            i = 1
            while i < len(media_queue):
                print(media_queue[i][0].chat.id)
                print(call.message.chat.id)
                if media_queue[i][0].chat.id == call.message.chat.id:


                    media_queue.pop(i)

                else:
                    i+=1
            bot.send_message(call.message.chat.id,
                     "Очередь удалена")
            
        elif call.data == "stop_transc_button":
            script_name = "diarize.py"
            result = subprocess.run(
                ['wmic', 'process', 'where', "name='python.exe'", 'get', 'ProcessId,CommandLine'],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True
                )
            for line in result.stdout.splitlines():
                if script_name in line:  # Ищем имя скрипта в строке
                    parts = line.split()  # Разбиваем строку на части
                    pid = parts[-1]  # Последний элемент — это PID
                    print(f"Найден процесс с PID {pid} для скрипта {script_name}")
                    
                    # Завершаем процесс
                    subprocess.run(['taskkill', '/PID', pid, '/F'], check=True)
                    bot.send_message(call.message.chat.id,
                     "Транскрибация прервана")
                    print(f"Процесс с PID {pid} завершен.")
                    break
                else:
                    print(f"Скрипт {script_name} не найден среди запущенных процессов.")

        



@bot.message_handler(content_types=['voice', 'audio', 'video', 'document'])
def get_audio_messages(message):  
    db_manager.make_new_session()

    keyboard = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    button_buy = telebot.types.KeyboardButton(text="Купить минуты")
    keyboard.add(button_buy)
    if message.content_type == "voice":
        media_file = message.voice
    elif message.content_type == 'video':
        media_file = message.video
    elif message.content_type == 'document':
        media_file = message.document
        a = media_file.file_name.split(".")
        if a[-1] == "wav":
            pass
        else:
            db_manager.close_session()
            return
    else:

        media_file = message.audio

    # duration = media_file.duration
    # if not db_manager.is_in_database(message.from_user.id):
    #         db_manager.add_user(message.from_user.id)

    # if duration > 120 and message.from_user.id != 197016529 and message.from_user.id != 5085651430 and message.from_user.id != 220183123:
    #     user  = db_manager.get_user(message.from_user.id)
    #     if user.paid:
    #         if user.minutes - duration > 0:
    #             process_audio(bot, message)
    #             db_manager.set_new_minutes(message.from_user.id, user.minutes - duration)
    #         else:
    #             # keyboard = types.InlineKeyboardMarkup()
    #             # button1 = types.InlineKeyboardButton(text="Купить минуты", callback_data="button1")
    #             # keyboard.add(button1)
    #             bot.send_message(message.chat.id, "У вас недостаточно минут на балансе",  reply_markup=keyboard)
    #     else:
    #         # keyboard = types.InlineKeyboardMarkup()
    #         # button1 = types.InlineKeyboardButton(text="Купить минуты", callback_data="button1")
    #         # keyboard.add(button1)
    #         bot.send_message(message.chat.id, "Бесплатный аккаунт может обработать аудио до 2 минут",  reply_markup=keyboard)
    # else:
    if message.content_type == "voice":
        media_file = message.voice
    elif message.content_type == 'video':
        media_file = message.video
    elif message.content_type == "document":
        media_file = message.document
    else:
        media_file = message.audio

    
    if message.media_group_id in media_groups_in_work.keys():
        media_groups_in_work[message.media_group_id].append(message)
    elif message.media_group_id:
        media_groups_in_work[message.media_group_id] = [message]

    if media_queue:
        insert_in_queue(media_queue, message, media_file.file_name)
    else:
        media_queue.append([message, media_file.file_name])
    bot.send_message(message.chat.id, "Файл добавлен в очередь")
    


    if not meida_queue_thread.is_alive():
        meida_queue_thread.start()
        # if result:
        #     if user.minutes - duration < 120: ошибка UnboundLocalError: cannot access local variable 'user' where it is not associated with a value
        #         db_manager.set_new_minutes(message.from_user.id, 0)
        #     else:
        #         db_manager.set_new_minutes(message.from_user.id, user.minutes - duration)
    db_manager.close_session()

# @bot.shipping_query_handler(func=lambda query: True)
# def shipping(shipping_query):
#     print(shipping_query)
#     bot.answer_shipping_query(shipping_query.id, ok=True, shipping_options=shipping_options,
#                               error_message='Oh, seems like our Dog couriers are having a lunch right now. Try again later!')


@bot.pre_checkout_query_handler(func=lambda query: True)
def checkout(pre_checkout_query):
    print(f"Received pre_checkout_query: {pre_checkout_query}")
    bot.answer_pre_checkout_query(pre_checkout_query.id, ok=True)




@bot.message_handler(content_types=['successful_payment'])
def got_payment(message):
    bot.send_message(message.chat.id,
                     'Оплата `{} {}` прошла успешно '
                     'Спасибо за покупку!'.format(
                         message.successful_payment.total_amount / 100, message.successful_payment.currency),
                     parse_mode='Markdown')
    db_manager.make_new_session()
    user  = db_manager.get_user(message.from_user.id)
    added_minutes = 0
    if user.paid:
        added_minutes = user.minutes
    if message.successful_payment.total_amount == 50000:
        db_manager.set_new_minutes(message.from_user.id, added_minutes + 3600)
    elif message.successful_payment.total_amount == 100000:
        db_manager.set_new_minutes(message.from_user.id, added_minutes + 18000)
    elif message.successful_payment.total_amount == 250000:
        db_manager.set_new_minutes(message.from_user.id, added_minutes + 36000)
    
    db_manager.close_session()
    

#bot.polling(none_stop=True, allowed_updates=[])     
bot.infinity_polling(timeout=10, long_polling_timeout = 5, allowed_updates=[])